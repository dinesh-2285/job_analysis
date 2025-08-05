# src/resume_processor_advanced.py

'''import pandas as pd
import numpy as np
import re
import string
import logging
from datetime import datetime
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

# NLP and ML imports
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.stem import WordNetLemmatizer
from nltk.chunk import ne_chunk
from nltk.tag import pos_tag

# PDF and document processing
import PyPDF2
import docx
from io import BytesIO
import fitz  # PyMuPDF for better PDF extraction

# Advanced NLP
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import MinMaxScaler

# Fuzzy matching
from fuzzywuzzy import fuzz, process

# Configuration
from job_analysis.config import PROCESSED_DATA_PATH, LOG_FILE

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
    
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')
    
try:
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('wordnet')

try:
    nltk.data.find('taggers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('averaged_perceptron_tagger')

try:
    nltk.data.find('chunkers/maxent_ne_chunker')
except LookupError:
    nltk.download('maxent_ne_chunker')

try:
    nltk.data.find('corpora/words')
except LookupError:
    nltk.download('words')

class AdvancedResumeProcessor:
    """
    Advanced Resume Processing and Job Matching Engine
    Uses NLP, semantic similarity, and intelligent extraction
    """
    
    def __init__(self):
        self.job_df = None
        self.sentence_model = None
        self.vectorizer = None
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words('english'))
        self.load_models()
        self.load_job_data()
        self.skill_synonyms = self.create_skill_synonyms()
        
    def load_models(self):
        """Load pre-trained models for semantic analysis"""
        try:
            print("🤖 Loading semantic similarity model...")
            # Load sentence transformer for semantic similarity
            self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
            print("✅ Semantic model loaded successfully")
            
        except Exception as e:
            logging.error(f"Error loading models: {str(e)}")
            print(f"⚠️ Warning: Could not load semantic model - {str(e)}")
            self.sentence_model = None
    
    def load_job_data(self):
        """Load job dataset for matching"""
        try:
            self.job_df = pd.read_csv(PROCESSED_DATA_PATH)
            print(f"✅ Loaded {len(self.job_df)} job records for matching")
            
        except Exception as e:
            logging.error(f"Error loading job data: {str(e)}")
            print(f"❌ Error loading job data: {str(e)}")
            raise
    
    def create_skill_synonyms(self):
        """Create skill synonyms dictionary for better matching"""
        return {
            'python': ['python', 'py', 'python3', 'python2'],
            'javascript': ['javascript', 'js', 'node.js', 'nodejs', 'react', 'angular', 'vue'],
            'java': ['java', 'j2ee', 'spring', 'hibernate'],
            'machine learning': ['ml', 'machine learning', 'artificial intelligence', 'ai', 'deep learning'],
            'data science': ['data science', 'data analysis', 'analytics', 'big data'],
            'sql': ['sql', 'mysql', 'postgresql', 'sqlite', 'database'],
            'aws': ['aws', 'amazon web services', 'ec2', 's3', 'lambda'],
            'docker': ['docker', 'containerization', 'kubernetes', 'k8s'],
            'git': ['git', 'github', 'version control', 'gitlab', 'bitbucket'],
            'agile': ['agile', 'scrum', 'kanban', 'sprint'],
            'project management': ['project management', 'pmp', 'pm', 'program management'],
            'communication': ['communication', 'interpersonal', 'teamwork', 'collaboration']
        }
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from PDF using PyMuPDF (better than PyPDF2)"""
        try:
            if isinstance(pdf_file, str):
                # File path
                doc = fitz.open(pdf_file)
            else:
                # File-like object
                doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
            
            text = ""
            for page in doc:
                text += page.get_text()
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            logging.error(f"Error extracting PDF text: {str(e)}")
            # Fallback to PyPDF2
            try:
                if isinstance(pdf_file, str):
                    with open(pdf_file, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                else:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                return text.strip()
            except Exception as e2:
                logging.error(f"Fallback PDF extraction failed: {str(e2)}")
                return ""
    
    def extract_text_from_docx(self, docx_file):
        """Extract text from DOCX file"""
        try:
            if isinstance(docx_file, str):
                doc = docx.Document(docx_file)
            else:
                doc = docx.Document(docx_file)
            
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            return '\n'.join(text)
            
        except Exception as e:
            logging.error(f"Error extracting DOCX text: {str(e)}")
            return ""
    
    def extract_text_from_file(self, file_path_or_object):
        """Extract text from various file formats"""
        try:
            if isinstance(file_path_or_object, str):
                file_path = Path(file_path_or_object)
                extension = file_path.suffix.lower()
                
                if extension == '.pdf':
                    return self.extract_text_from_pdf(file_path_or_object)
                elif extension == '.docx':
                    return self.extract_text_from_docx(file_path_or_object)
                elif extension == '.txt':
                    with open(file_path_or_object, 'r', encoding='utf-8') as f:
                        return f.read()
                else:
                    return ""
            else:
                # Handle uploaded file objects
                file_name = getattr(file_path_or_object, 'name', '')
                if file_name.endswith('.pdf'):
                    return self.extract_text_from_pdf(file_path_or_object)
                elif file_name.endswith('.docx'):
                    return self.extract_text_from_docx(file_path_or_object)
                elif file_name.endswith('.txt'):
                    return file_path_or_object.read().decode('utf-8')
                else:
                    return ""
                    
        except Exception as e:
            logging.error(f"Error extracting text from file: {str(e)}")
            return ""
    
    def clean_text(self, text):
        """Clean and preprocess text"""
        try:
            # Remove extra whitespace and normalize
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            # Remove special characters but keep some punctuation
            text = re.sub(r'[^\w\s\-\.\,\(\)]', ' ', text)
            
            # Remove extra spaces
            text = re.sub(r'\s+', ' ', text)
            
            return text
            
        except Exception as e:
            logging.error(f"Error cleaning text: {str(e)}")
            return text
    
    def extract_contact_info(self, text):
        """Extract contact information from resume text"""
        contact_info = {}
        
        try:
            # Email extraction
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            emails = re.findall(email_pattern, text)
            contact_info['emails'] = emails
            
            # Phone number extraction
            phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
            phones = re.findall(phone_pattern, text)
            contact_info['phones'] = phones
            
            # LinkedIn profile extraction
            linkedin_pattern = r'(https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+'
            linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
            contact_info['linkedin'] = linkedin
            
            # GitHub profile extraction
            github_pattern = r'(https?://)?(?:www\.)?github\.com/[\w\-]+'
            github = re.findall(github_pattern, text, re.IGNORECASE)
            contact_info['github'] = github
            
        except Exception as e:
            logging.error(f"Error extracting contact info: {str(e)}")
        
        return contact_info
    
    def extract_education(self, text):
        """Extract education information"""
        education = []
        
        try:
            # Common degree patterns
            degree_patterns = [
                r'(Bachelor|Master|PhD|Ph\.D|M\.S|B\.S|B\.A|M\.A|MBA|M\.Tech|B\.Tech)\s+(?:of\s+)?(?:Science\s+)?(?:Arts\s+)?(?:in\s+)?([A-Za-z\s]+)',
                r'(B\.E|M\.E|BE|ME)\s+(?:in\s+)?([A-Za-z\s]+)',
                r'(Diploma|Certificate)\s+(?:in\s+)?([A-Za-z\s]+)'
            ]
            
            for pattern in degree_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                for match in matches:
                    education.append({
                        'degree': match[0],
                        'field': match[1].strip()
                    })
            
            # University/College names
            university_pattern = r'(University|College|Institute|School)\s+of\s+([A-Za-z\s]+)'
            universities = re.findall(university_pattern, text, re.IGNORECASE)
            
            return {
                'degrees': education,
                'institutions': universities
            }
            
        except Exception as e:
            logging.error(f"Error extracting education: {str(e)}")
            return {'degrees': [], 'institutions': []}
    
    def extract_experience(self, text):
        """Extract work experience information"""
        experience = []
        
        try:
            # Year patterns
            year_pattern = r'(19|20)\d{2}'
            years = re.findall(year_pattern, text)
            
            # Company patterns (look for common indicators)
            company_indicators = ['worked at', 'employed by', 'company:', 'organization:']
            
            # Job title patterns
            title_patterns = [
                r'(Software Engineer|Developer|Analyst|Manager|Director|Senior|Junior|Lead|Principal)\s+([A-Za-z\s]+)',
                r'(Data Scientist|Product Manager|Project Manager|Team Lead|Technical Lead)'
            ]
            
            job_titles = []
            for pattern in title_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                job_titles.extend(matches)
            
            return {
                'years_mentioned': years,
                'job_titles': job_titles,
                'total_experience_years': len(set(years))
            }
            
        except Exception as e:
            logging.error(f"Error extracting experience: {str(e)}")
            return {'years_mentioned': [], 'job_titles': [], 'total_experience_years': 0}
    
    def extract_skills_advanced(self, text):
        """Advanced skill extraction with NLP and pattern matching"""
        skills = set()
        
        try:
            # Convert to lowercase for processing
            text_lower = text.lower()
            
            # Technical skills database
            technical_skills = {
                'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin'],
                'web': ['html', 'css', 'react', 'angular', 'vue', 'nodejs', 'express', 'django', 'flask', 'laravel'],
                'database': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle', 'sqlite'],
                'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'ansible'],
                'data': ['pandas', 'numpy', 'scipy', 'matplotlib', 'seaborn', 'plotly', 'tableau', 'power bi'],
                'ml': ['scikit-learn', 'tensorflow', 'pytorch', 'keras', 'opencv', 'nltk', 'spacy'],
                'tools': ['git', 'jenkins', 'jira', 'confluence', 'slack', 'trello', 'figma', 'sketch'],
                'methodologies': ['agile', 'scrum', 'kanban', 'devops', 'ci/cd', 'tdd', 'bdd']
            }
            
            # Extract skills from predefined lists
            for category, skill_list in technical_skills.items():
                for skill in skill_list:
                    if skill in text_lower:
                        skills.add(skill)
            
            # Extract skills using pattern matching
            skill_patterns = [
                r'skills?:?\s*([A-Za-z0-9\s,\-\.\+#]+)',
                r'technologies?:?\s*([A-Za-z0-9\s,\-\.\+#]+)',
                r'tools?:?\s*([A-Za-z0-9\s,\-\.\+#]+)',
                r'languages?:?\s*([A-Za-z0-9\s,\-\.\+#]+)'
            ]
            
            for pattern in skill_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                for match in matches:
                    # Split by common separators
                    skill_parts = re.split(r'[,;|\n]', match)
                    for part in skill_parts:
                        cleaned_skill = part.strip().lower()
                        if len(cleaned_skill) > 1 and cleaned_skill not in self.stop_words:
                            skills.add(cleaned_skill)
            
            # Use NLP to extract potential skills
            tokens = word_tokenize(text_lower)
            pos_tags = pos_tag(tokens)
            
            # Extract nouns that might be skills
            for token, pos in pos_tags:
                if pos in ['NN', 'NNS'] and len(token) > 2:
                    if token not in self.stop_words and token.isalpha():
                        # Check if it's a known skill
                        for skill_list in technical_skills.values():
                            if token in skill_list:
                                skills.add(token)
            
            return list(skills)
            
        except Exception as e:
            logging.error(f"Error extracting skills: {str(e)}")
            return []
    
    def normalize_skills(self, skills):
        """Normalize skills using synonyms"""
        normalized = set()
        
        for skill in skills:
            skill_lower = skill.lower().strip()
            added = False
            
            # Check against synonyms
            for canonical_skill, synonyms in self.skill_synonyms.items():
                if skill_lower in synonyms:
                    normalized.add(canonical_skill)
                    added = True
                    break
            
            # If not found in synonyms, add as is
            if not added:
                normalized.add(skill_lower)
        
        return list(normalized)
    
    def analyze_resume(self, file_path_or_object):
        """Complete resume analysis pipeline"""
        try:
            print("📄 Analyzing resume...")
            
            # Extract text
            text = self.extract_text_from_file(file_path_or_object)
            if not text:
                return {'error': 'Could not extract text from file'}
            
            # Clean text
            clean_text = self.clean_text(text)
            
            # Extract information
            contact_info = self.extract_contact_info(text)
            education = self.extract_education(text)
            experience = self.extract_experience(text)
            skills = self.extract_skills_advanced(text)
            normalized_skills = self.normalize_skills(skills)
            
            print(f"✅ Resume analysis complete - found {len(normalized_skills)} skills")
            
            return {
                'raw_text': text,
                'clean_text': clean_text,
                'contact_info': contact_info,
                'education': education,
                'experience': experience,
                'skills': normalized_skills,
                'word_count': len(clean_text.split()),
                'analysis_timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logging.error(f"Error analyzing resume: {str(e)}")
            return {'error': str(e)}
    
    def calculate_skill_match_score(self, resume_skills, job_skills):
        """Calculate skill match score between resume and job"""
        try:
            if not resume_skills or not job_skills:
                return 0.0
            
            # Convert to sets for easier comparison
            resume_set = set([skill.lower() for skill in resume_skills])
            job_set = set([skill.lower().strip() for skill in job_skills.split(',')])
            
            # Direct matches
            direct_matches = resume_set.intersection(job_set)
            
            # Fuzzy matches
            fuzzy_matches = 0
            for resume_skill in resume_set:
                for job_skill in job_set:
                    if fuzz.ratio(resume_skill, job_skill) > 80:
                        fuzzy_matches += 1
                        break
            
            # Synonym matches
            synonym_matches = 0
            for resume_skill in resume_set:
                for canonical_skill, synonyms in self.skill_synonyms.items():
                    if resume_skill in synonyms:
                        for job_skill in job_set:
                            if job_skill in synonyms:
                                synonym_matches += 1
                                break
            
            # Calculate weighted score
            total_matches = len(direct_matches) + (fuzzy_matches * 0.8) + (synonym_matches * 0.6)
            total_job_skills = len(job_set)
            
            if total_job_skills == 0:
                return 0.0
            
            score = min(100.0, (total_matches / total_job_skills) * 100)
            return score
            
        except Exception as e:
            logging.error(f"Error calculating skill match score: {str(e)}")
            return 0.0
    
    def calculate_semantic_similarity(self, resume_text, job_description):
        """Calculate semantic similarity using sentence transformers"""
        try:
            if not self.sentence_model:
                return 0.0
            
            # Encode texts
            resume_embedding = self.sentence_model.encode([resume_text])
            job_embedding = self.sentence_model.encode([job_description])
            
            # Calculate cosine similarity
            similarity = cosine_similarity(resume_embedding, job_embedding)[0][0]
            return float(similarity * 100)  # Convert to percentage
            
        except Exception as e:
            logging.error(f"Error calculating semantic similarity: {str(e)}")
            return 0.0
    
    def find_matching_jobs(self, resume_analysis, top_n=10):
        """Find best matching jobs for a resume"""
        try:
            if 'error' in resume_analysis:
                return {'error': resume_analysis['error']}
            
            print(f"🔍 Finding top {top_n} matching jobs...")
            
            resume_skills = resume_analysis['skills']
            resume_text = resume_analysis['clean_text']
            
            matches = []
            
            for idx, job in self.job_df.iterrows():
                # Calculate skill match score
                skill_score = self.calculate_skill_match_score(resume_skills, job['Skills'])
                
                # Calculate semantic similarity
                job_description = f"{job['Job Title']} {job['Skills']}"
                semantic_score = self.calculate_semantic_similarity(resume_text, job_description)
                
                # Combined score (weighted average)
                combined_score = (skill_score * 0.6) + (semantic_score * 0.4)
                
                matches.append({
                    'job_id': idx,
                    'job_title': job['Job Title'],
                    'company': job['Company'],
                    'location': job['Location'],
                    'stream': job['Stream'],
                    'required_skills': job['Skills'],
                    'skill_match_score': skill_score,
                    'semantic_similarity': semantic_score,
                    'combined_score': combined_score,
                    'matching_skills': list(set(resume_skills).intersection(
                        set([skill.lower().strip() for skill in job['Skills'].split(',')])
                    ))
                })
            
            # Sort by combined score
            matches.sort(key=lambda x: x['combined_score'], reverse=True)
            
            print(f"✅ Found {len(matches)} job matches")
            
            return {
                'total_matches': len(matches),
                'top_matches': matches[:top_n],
                'resume_skills': resume_skills,
                'analysis_timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logging.error(f"Error finding matching jobs: {str(e)}")
            return {'error': str(e)}
    
    def generate_improvement_suggestions(self, resume_analysis, job_matches):
        """Generate suggestions for resume improvement"""
        try:
            if 'error' in resume_analysis or 'error' in job_matches:
                return {'error': 'Cannot generate suggestions due to analysis errors'}
            
            suggestions = []
            resume_skills = set(resume_analysis['skills'])
            
            # Analyze top job requirements
            top_jobs = job_matches['top_matches'][:5]
            required_skills = []
            
            for job in top_jobs:
                job_skills = [skill.lower().strip() for skill in job['required_skills'].split(',')]
                required_skills.extend(job_skills)
            
            # Find most common missing skills
            skill_frequency = {}
            for skill in required_skills:
                if skill not in resume_skills:
                    skill_frequency[skill] = skill_frequency.get(skill, 0) + 1
            
            # Sort by frequency
            missing_skills = sorted(skill_frequency.items(), key=lambda x: x[1], reverse=True)
            
            # Generate suggestions
            if missing_skills:
                suggestions.append({
                    'category': 'Skills Gap',
                    'priority': 'High',
                    'suggestion': f"Consider adding these in-demand skills: {', '.join([skill for skill, _ in missing_skills[:5]])}",
                    'impact': 'Adding these skills could significantly improve your job match scores'
                })
            
            # Experience suggestions
            if resume_analysis['experience']['total_experience_years'] < 2:
                suggestions.append({
                    'category': 'Experience',
                    'priority': 'Medium',
                    'suggestion': 'Consider adding more details about your projects, internships, or volunteer work',
                    'impact': 'More experience details can improve your profile completeness'
                })
            
            # Education suggestions
            if not resume_analysis['education']['degrees']:
                suggestions.append({
                    'category': 'Education',
                    'priority': 'Medium',
                    'suggestion': 'Ensure your educational background is clearly mentioned',
                    'impact': 'Clear education details help with initial screening'
                })
            
            # Contact info suggestions
            if not resume_analysis['contact_info']['emails']:
                suggestions.append({
                    'category': 'Contact Information',
                    'priority': 'High',
                    'suggestion': 'Add your email address and phone number',
                    'impact': 'Essential for recruiters to contact you'
                })
            
            return {
                'suggestions': suggestions,
                'missing_skills': missing_skills[:10],
                'skill_gap_analysis': {
                    'resume_skills_count': len(resume_skills),
                    'common_job_skills': len(set(required_skills)),
                    'overlap_percentage': (len(resume_skills.intersection(set(required_skills))) / len(set(required_skills))) * 100 if required_skills else 0
                }
            }
            
        except Exception as e:
            logging.error(f"Error generating improvement suggestions: {str(e)}")
            return {'error': str(e)}

# Usage example and testing
if __name__ == "__main__":
    # Initialize processor
    processor = AdvancedResumeProcessor()
    
    # Test with sample resume (you would provide actual file path)
    print("🚀 Advanced Resume Processor ready for use!")
    print("Features available:")
    print("  ✅ Multi-format support (PDF, DOCX, TXT)")
    print("  ✅ Advanced NLP skill extraction")
    print("  ✅ Semantic similarity matching")
    print("  ✅ Intelligent job recommendations")
    print("  ✅ Resume improvement suggestions")'''
# src/resume_processor_advanced.py

import pandas as pd
import numpy as np
import re
import logging
from datetime import datetime
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

# NLP and ML imports
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# PDF and document processing
import fitz  # PyMuPDF
import docx

# Configuration
from job_analysis.config import PROCESSED_DATA_PATH

class AdvancedResumeProcessor:
    """
    Advanced Resume Processing and Job Matching Engine
    Uses NLP, semantic similarity, and intelligent extraction
    """
    def __init__(self):
        self.job_df = None
        self.sentence_model = None
        self.job_embeddings = None
        self.job_skills_sets = None
        self.SKILLS_DB = [
            'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin', 'typescript',
            'html', 'css', 'react', 'angular', 'vue', 'nodejs', 'express', 'django', 'flask', 'laravel', 'spring boot',
            'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle', 'sqlite', 'nosql',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'ansible', 'ci/cd', 'jenkins', 'github actions',
            'pandas', 'numpy', 'scipy', 'matplotlib', 'seaborn', 'plotly', 'tableau', 'power bi', 'd3.js',
            'scikit-learn', 'tensorflow', 'pytorch', 'keras', 'opencv', 'nltk', 'spacy', 'hugging face', 'llm',
            'machine learning', 'data science', 'deep learning', 'artificial intelligence', 'nlp', 'computer vision',
            'data analysis', 'data visualization', 'statistical analysis', 'data modeling', 'etl',
            'git', 'jira', 'confluence', 'slack', 'trello', 'figma', 'sketch', 'adobe xd',
            'agile', 'scrum', 'kanban', 'devops', 'tdd', 'bdd', 'project management', 'product management'
        ]
        self.load_models()
        self.load_job_data_and_prepare_embeddings()

    def load_models(self):
        try:
            self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
        except Exception as e:
            logging.error(f"Error loading models: {str(e)}")
            raise

    def load_job_data_and_prepare_embeddings(self):
        try:
            self.job_df = pd.read_csv(PROCESSED_DATA_PATH)
            self.job_df.dropna(subset=['Job Title', 'Skills'], inplace=True)
            if self.sentence_model:
                job_descriptions = (self.job_df['Job Title'] + ' ' + self.job_df['Skills']).tolist()
                self.job_embeddings = self.sentence_model.encode(job_descriptions, show_progress_bar=True)
            self.job_skills_sets = [set(str(s).lower().strip().split(',')) for s in self.job_df['Skills']]
        except Exception as e:
            logging.error(f"Error loading job data: {str(e)}")
            raise

    def extract_text_from_file(self, file_object):
        try:
            if file_object.name.endswith('.pdf'):
                doc = fitz.open(stream=file_object.read(), filetype="pdf")
                text = "".join(page.get_text() for page in doc)
                doc.close()
                return text
            elif file_object.name.endswith('.docx'):
                return "\n".join([p.text for p in docx.Document(file_object).paragraphs])
            return ""
        except Exception as e:
            logging.error(f"Error extracting text: {e}")
            return ""

    def extract_contact_info(self, text):
        info = {}
        info['emails'] = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        info['phones'] = re.findall(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
        return info

    def extract_skills_advanced(self, text):
        skills = set()
        text_lower = text.lower()
        for skill in self.SKILLS_DB:
            if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                skills.add(skill)
        return list(skills)

    def extract_education(self, text):
        education = []
        keywords = ['B.Tech', 'M.Tech', 'B.E', 'M.E', 'B.Sc', 'M.Sc', 'B.A', 'M.A', 'Bachelor', 'Master', 'MBA', 'PhD', 'Intermediate']
        pattern = re.compile(r'(?i)(' + '|'.join(keywords).replace('.', r'\.') + r')\b(.*)')
        for line in text.split('\n'):
            match = pattern.search(line)
            if match:
                degree, field = match.group(1).strip(), match.group(2).strip()
                # Corrected line: removed the inline flag (?i)
                cleaned_field = re.sub(r'^\s*(?:in|of|from|,|:)\s*', '', field, flags=re.IGNORECASE).strip()
                if cleaned_field:
                    education.append({'degree': degree, 'field': cleaned_field})
        return {'degrees': education}

    def extract_experience(self, text):
        years = re.findall(r'\b(19|20)\d{2}\b', text)
        return {'total_experience_years': len(set(years)) if years else 1}

    def analyze_resume(self, file_object):
        text = self.extract_text_from_file(file_object)
        if not text: return {'error': 'Could not extract text'}
        clean_text = ' '.join(text.split())
        return {
            'clean_text': clean_text,
            'contact_info': self.extract_contact_info(clean_text),
            'education': self.extract_education(text), # Use original text for line-based parsing
            'experience': self.extract_experience(clean_text),
            'skills': self.extract_skills_advanced(clean_text),
            'word_count': len(clean_text.split()),
        }

    def find_matching_jobs(self, resume_analysis, top_n=20):
        if 'error' in resume_analysis: return {'error': resume_analysis['error']}
        resume_skills_set = set(resume_analysis['skills'])
        skill_scores = np.array([(len(resume_skills_set.intersection(js)) / len(js) * 100 if js else 0) for js in self.job_skills_sets])
        if self.job_embeddings is not None:
            resume_embedding = self.sentence_model.encode([resume_analysis['clean_text']])
            semantic_scores = cosine_similarity(resume_embedding, self.job_embeddings)[0] * 100
        else:
            semantic_scores = np.zeros(len(self.job_df))
        combined_scores = (skill_scores * 0.6) + (semantic_scores * 0.4)
        top_indices = np.argsort(combined_scores)[-top_n:][::-1]
        matches = []
        for i in top_indices:
            job = self.job_df.iloc[i]
            matches.append({
                'job_title': job['Job Title'], 'company': job['Company'],
                'location': job['Location'], 'stream': job['Stream'], 'required_skills': job['Skills'],
                'skill_match_score': skill_scores[i], 'semantic_similarity': semantic_scores[i],
                'combined_score': combined_scores[i],
                'matching_skills': list(resume_skills_set.intersection(self.job_skills_sets[i]))
            })
        return {'top_matches': matches, 'resume_skills': list(resume_skills_set)}

    def generate_improvement_suggestions(self, resume_analysis, job_matches):
        if 'error' in resume_analysis or 'error' in job_matches:
            return {'error': 'Cannot generate suggestions'}
        suggestions = []
        resume_skills = set(resume_analysis.get('skills', []))
        top_jobs = job_matches.get('top_matches', [])[:5]
        if not top_jobs:
            return {'suggestions': [{'category': 'General', 'priority': 'Low', 'suggestion': 'Could not find job matches to compare against.', 'impact': 'Suggestions are based on job data.'}]}
        required_skills = [s.strip().lower() for j in top_jobs for s in str(j.get('required_skills', '')).split(',') if s]
        skill_freq = {s: required_skills.count(s) for s in set(required_skills) if s not in resume_skills}
        missing_skills = sorted(skill_freq.items(), key=lambda x: x[1], reverse=True)
        if missing_skills:
            suggestions.append({
                'category': 'Skills Gap', 'priority': 'High',
                'suggestion': f"Consider adding these in-demand skills: {', '.join([s for s, _ in missing_skills[:5]])}",
                'impact': 'Improves your match score.'
            })
        if not resume_analysis.get('education', {}).get('degrees'):
            suggestions.append({
                'category': 'Education', 'priority': 'Medium',
                'suggestion': 'Ensure your educational background is clearly listed.',
                'impact': 'A key screening factor.'
            })
        return {
            'suggestions': suggestions, 'missing_skills': missing_skills[:10],
            'skill_gap_analysis': {
                'resume_skills_count': len(resume_skills),
                'common_job_skills': len(set(required_skills)),
                'overlap_percentage': (len(resume_skills.intersection(set(required_skills))) / len(set(required_skills))) * 100 if required_skills else 0
            }
        }
